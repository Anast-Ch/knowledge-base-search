import io
import os
from fastapi import HTTPException
import pdfplumber
from docx import Document

async def parse_pdf(file_bytes: bytes) -> str:
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if len(pdf.pages) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="PDF файл пуст или не содержит страниц"
                )
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="PDF файл не содержит текста (возможно, это сканированный документ или изображение)"
            )
        
        return text
        
    except HTTPException:
        raise
    except pdfplumber.exceptions.PDFSyntaxError as e:
        raise HTTPException(
            status_code=400,
            detail=f"Поврежденный PDF файл: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Внутренняя ошибка при парсинге PDF: {str(e)}"
        )


async def parse_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        if len(doc.paragraphs) == 0:
            raise HTTPException(
                status_code=400,
                detail="DOCX файл пуст или не содержит текста"
            )
        
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="DOCX файл не содержит текста"
            )
        
        return text
        
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e).lower()
        if any(keyword in error_msg for keyword in [
            "zip", "invalid", "corrupt", "not a zip", "bad file"
        ]):
            raise HTTPException(
                status_code=400,
                detail=f"Поврежденный DOCX файл: {str(e)}"
            )
        
        raise HTTPException(
            status_code=500,
            detail=f"Внутренняя ошибка при парсинге DOCX: {str(e)}"
        )


async def parse_document(file_bytes: bytes, filename: str) -> str:
    if not file_bytes or len(file_bytes) == 0:
        raise HTTPException(
            status_code=400,
            detail="Загружен пустой файл"
        )
    
    if len(file_bytes) < 100:  # 100 байт
        raise HTTPException(
            status_code=400,
            detail="Файл слишком маленький (менее 100 байт)"
        )
    
    file_extension = filename.lower().split('.')[-1] if '.' in filename else ''
    
    if file_extension == 'pdf':
        return await parse_pdf(file_bytes)
    elif file_extension == 'docx':
        return await parse_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Неподдерживаемый формат файла: {file_extension}"
        )